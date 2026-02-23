"""Document parser -- converts files into structured sections.

Supports: PDF, DOCX, HTML, Markdown, CSV, JSON, plain text.

Security: All file access goes through validate_file_for_ingestion()
which prevents path traversal, enforces size limits, and validates
file extensions.
"""

import csv
import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

from bitmod.security import MAX_FILE_SIZE_BYTES, validate_file_for_ingestion

logger = logging.getLogger(__name__)

# Maximum rows to parse from CSV to prevent memory exhaustion
MAX_CSV_ROWS = 100_000

# Maximum JSON file size for in-memory parsing (10 MB)
MAX_JSON_PARSE_SIZE = 10 * 1024 * 1024


@dataclass
class ParsedSection:
    title: str
    text: str
    section_number: str | None = None
    hierarchy_path: list[str] | None = None
    metadata: dict | None = None


@dataclass
class ParsedDocument:
    title: str
    sections: list[ParsedSection]
    source_format: str
    metadata: dict = field(default_factory=dict)


def parse_file(
    file_path: str,
    title: str | None = None,
    allowed_base_dirs: list[str] | None = None,
) -> ParsedDocument:
    """Parse a file into structured sections based on its format.

    Args:
        file_path: Path to the file to parse.
        title: Optional title override.
        allowed_base_dirs: Optional list of allowed base directories.
            If provided, the file must be under one of these directories.

    Raises:
        ValueError: If file fails validation (bad extension, too large, traversal).
        FileNotFoundError: If file does not exist.
    """
    # Validate path security: traversal, extension, size
    resolved_path = validate_file_for_ingestion(file_path, allowed_base_dirs)
    path = Path(resolved_path)

    suffix = path.suffix.lower()
    title = title or path.stem

    if suffix == ".pdf":
        return _parse_pdf(path, title)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(path, title)
    elif suffix in (".html", ".htm"):
        return _parse_html(path, title)
    elif suffix in (".md", ".markdown"):
        return _parse_markdown(path, title)
    elif suffix == ".csv":
        return _parse_csv(path, title)
    elif suffix == ".json":
        return _parse_json(path, title)
    elif suffix in (".txt", ".text", ".log", ".rst"):
        return _parse_text_file(path, title)
    else:
        # Fall back to plain text for allowed extensions
        return _parse_text_file(path, title)


def parse_text(
    text: str,
    title: str = "Untitled",
    source_format: str = "text",
) -> ParsedDocument:
    """Parse raw text into sections by splitting on blank lines or headings."""
    if not isinstance(text, str):
        raise ValueError("Input text must be a string")
    # Enforce size limit on raw text input too
    if len(text) > MAX_FILE_SIZE_BYTES:
        raise ValueError(f"Text input exceeds maximum size ({MAX_FILE_SIZE_BYTES:,} bytes)")
    sections = _split_by_paragraphs(text)
    return ParsedDocument(title=title, sections=sections, source_format=source_format)


# ---------------------------------------------------------------------------
# Format-specific parsers
# ---------------------------------------------------------------------------


def _parse_pdf(path: Path, title: str) -> ParsedDocument:
    """Parse PDF using PyMuPDF (fitz) or pdfplumber.

    Note: PDF parsing libraries extract text only. Embedded scripts,
    JavaScript actions, and form submissions in malicious PDFs are not
    executed because we only call get_text() / extract_text().
    """
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(path))
        sections = []
        for i, page in enumerate(doc):
            text = page.get_text("text").strip()
            if text:
                sections.append(
                    ParsedSection(
                        title=f"Page {i + 1}",
                        text=text,
                        section_number=str(i + 1),
                        hierarchy_path=[title, f"Page {i + 1}"],
                    )
                )
        doc.close()
        return ParsedDocument(title=title, sections=sections, source_format="pdf")
    except ImportError:
        pass

    try:
        import pdfplumber

        sections = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                text = text.strip()
                if text:
                    sections.append(
                        ParsedSection(
                            title=f"Page {i + 1}",
                            text=text,
                            section_number=str(i + 1),
                            hierarchy_path=[title, f"Page {i + 1}"],
                        )
                    )
        return ParsedDocument(title=title, sections=sections, source_format="pdf")
    except ImportError:
        raise ImportError(
            "PDF parsing requires PyMuPDF or pdfplumber. Install: pip install pymupdf  or  pip install pdfplumber"
        )


def _parse_docx(path: Path, title: str) -> ParsedDocument:
    """Parse DOCX by extracting paragraphs grouped by headings.

    Note: python-docx only reads text and paragraph styles.
    Embedded macros (VBA) are not executed.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("DOCX parsing requires python-docx. Install: pip install python-docx")

    doc = DocxDocument(str(path))
    sections: list[ParsedSection] = []
    current_title = title
    current_text: list[str] = []
    section_num = 0

    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            # Save previous section
            if current_text:
                section_num += 1
                sections.append(
                    ParsedSection(
                        title=current_title,
                        text="\n".join(current_text),
                        section_number=str(section_num),
                        hierarchy_path=[title, current_title],
                    )
                )
                current_text = []
            current_title = para.text.strip() or f"Section {section_num + 1}"
        else:
            text = para.text.strip()
            if text:
                current_text.append(text)

    # Save last section
    if current_text:
        section_num += 1
        sections.append(
            ParsedSection(
                title=current_title,
                text="\n".join(current_text),
                section_number=str(section_num),
                hierarchy_path=[title, current_title],
            )
        )

    if not sections:
        # No headings found -- treat entire doc as one section
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        sections = [ParsedSection(title=title, text=full_text)]

    return ParsedDocument(title=title, sections=sections, source_format="docx")


def _parse_html(path: Path, title: str) -> ParsedDocument:
    """Parse HTML by extracting text grouped by heading tags.

    Uses html.parser backend (not lxml) to avoid processing instructions.
    Script and style tags are stripped.
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("HTML parsing requires BeautifulSoup. Install: pip install beautifulsoup4")

    with open(path, encoding="utf-8", errors="replace") as f:
        raw_html = f.read()

    # Use html.parser (pure Python, no XXE risk) instead of lxml
    soup = BeautifulSoup(raw_html, "html.parser")

    # Remove script and style elements to prevent embedded code leakage
    for tag in soup.find_all(["script", "style", "iframe", "object", "embed", "applet"]):
        tag.decompose()

    # Try to get title from <title> tag
    if soup.title and soup.title.string:
        title = soup.title.string.strip() or title

    sections: list[ParsedSection] = []
    heading_tags = {"h1", "h2", "h3", "h4", "h5", "h6"}
    current_title = title
    current_text: list[str] = []
    section_num = 0

    for element in soup.find_all(True):
        if element.name in heading_tags:
            if current_text:
                section_num += 1
                sections.append(
                    ParsedSection(
                        title=current_title,
                        text="\n".join(current_text),
                        section_number=str(section_num),
                    )
                )
                current_text = []
            current_title = element.get_text(strip=True)
        elif element.name in ("p", "li", "td", "pre", "blockquote"):
            text = element.get_text(strip=True)
            if text and len(text) > 10:
                current_text.append(text)

    if current_text:
        section_num += 1
        sections.append(
            ParsedSection(
                title=current_title,
                text="\n".join(current_text),
                section_number=str(section_num),
            )
        )

    if not sections:
        full_text = soup.get_text(separator="\n", strip=True)
        sections = [ParsedSection(title=title, text=full_text)]

    return ParsedDocument(title=title, sections=sections, source_format="html")


def _parse_markdown(path: Path, title: str) -> ParsedDocument:
    """Parse Markdown by splitting on heading lines (# ## ### etc.)."""
    with open(path, encoding="utf-8", errors="replace") as f:
        content = f.read()

    sections = _split_by_markdown_headings(content, title)
    return ParsedDocument(title=title, sections=sections, source_format="markdown")


def _parse_csv(path: Path, title: str) -> ParsedDocument:
    """Parse CSV -- each row group becomes a section.

    Enforces MAX_CSV_ROWS to prevent memory exhaustion from large files.
    """
    with open(path, encoding="utf-8", errors="replace") as f:
        reader = csv.DictReader(f)
        rows = []
        for i, row in enumerate(reader):
            if i >= MAX_CSV_ROWS:
                logger.warning("CSV truncated at %d rows (max: %d)", MAX_CSV_ROWS, MAX_CSV_ROWS)
                break
            rows.append(row)

    if not rows:
        return ParsedDocument(title=title, sections=[], source_format="csv")

    # Group rows into sections of 50 for reasonable chunk sizes
    group_size = 50
    sections = []
    for i in range(0, len(rows), group_size):
        group = rows[i : i + group_size]
        text = "\n".join(" | ".join(f"{k}: {v}" for k, v in row.items() if v) for row in group)
        sections.append(
            ParsedSection(
                title=f"Rows {i + 1}-{min(i + group_size, len(rows))}",
                text=text,
                section_number=str(i // group_size + 1),
                metadata={"row_start": i + 1, "row_end": min(i + group_size, len(rows))},
            )
        )

    return ParsedDocument(
        title=title,
        sections=sections,
        source_format="csv",
        metadata={"total_rows": len(rows), "columns": list(rows[0].keys())},
    )


def _parse_json(path: Path, title: str) -> ParsedDocument:
    """Parse JSON -- handles arrays of objects or nested structures.

    Enforces size limit to prevent memory exhaustion.
    """
    file_size = path.stat().st_size
    if file_size > MAX_JSON_PARSE_SIZE:
        raise ValueError(
            f"JSON file too large for in-memory parsing ({file_size:,} bytes, max {MAX_JSON_PARSE_SIZE:,})"
        )

    with open(path, encoding="utf-8", errors="replace") as f:
        data = json.load(f)

    sections = []
    if isinstance(data, list):
        group_size = 20
        for i in range(0, len(data), group_size):
            group = data[i : i + group_size]
            text = json.dumps(group, indent=2, default=str)
            sections.append(
                ParsedSection(
                    title=f"Items {i + 1}-{min(i + group_size, len(data))}",
                    text=text,
                    section_number=str(i // group_size + 1),
                )
            )
    elif isinstance(data, dict):
        for key, value in data.items():
            text = json.dumps(value, indent=2, default=str) if not isinstance(value, str) else value
            sections.append(ParsedSection(title=str(key), text=text))
    else:
        sections = [ParsedSection(title=title, text=str(data))]

    return ParsedDocument(title=title, sections=sections, source_format="json")


def _parse_text_file(path: Path, title: str) -> ParsedDocument:
    """Parse plain text by splitting on blank lines."""
    with open(path, encoding="utf-8", errors="replace") as f:
        content = f.read()

    sections = _split_by_paragraphs(content)
    return ParsedDocument(title=title, sections=sections, source_format="text")


# ---------------------------------------------------------------------------
# Text splitting helpers
# ---------------------------------------------------------------------------


def _split_by_paragraphs(text: str) -> list[ParsedSection]:
    """Split text into sections by double newlines (paragraphs)."""
    paragraphs = re.split(r"\n\s*\n", text.strip())
    sections = []
    current_text: list[str] = []
    section_num = 0

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        current_text.append(para)
        # Group paragraphs into ~2000 char sections
        if sum(len(t) for t in current_text) > 2000:
            section_num += 1
            sections.append(
                ParsedSection(
                    title=f"Section {section_num}",
                    text="\n\n".join(current_text),
                    section_number=str(section_num),
                )
            )
            current_text = []

    if current_text:
        section_num += 1
        sections.append(
            ParsedSection(
                title=f"Section {section_num}",
                text="\n\n".join(current_text),
                section_number=str(section_num),
            )
        )

    return sections or [ParsedSection(title="Content", text=text.strip())]


def _split_by_markdown_headings(text: str, doc_title: str) -> list[ParsedSection]:
    """Split markdown by heading lines."""
    lines = text.split("\n")
    sections: list[ParsedSection] = []
    current_title = doc_title
    current_lines: list[str] = []
    section_num = 0

    for line in lines:
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            if current_lines:
                section_num += 1
                sections.append(
                    ParsedSection(
                        title=current_title,
                        text="\n".join(current_lines).strip(),
                        section_number=str(section_num),
                        hierarchy_path=[doc_title, current_title],
                    )
                )
                current_lines = []
            current_title = heading_match.group(2).strip()
        else:
            current_lines.append(line)

    if current_lines:
        text_content = "\n".join(current_lines).strip()
        if text_content:
            section_num += 1
            sections.append(
                ParsedSection(
                    title=current_title,
                    text=text_content,
                    section_number=str(section_num),
                    hierarchy_path=[doc_title, current_title],
                )
            )

    return sections or [ParsedSection(title=doc_title, text=text.strip())]
